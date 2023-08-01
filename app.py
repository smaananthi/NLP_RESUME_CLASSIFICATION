import streamlit as st
from docx import Document
from Main_Module import *
import joblib
import pandas as pd


# Load the pre-trained model
model = joblib.load('all_models.pkl')

vectorizer = model['text_vectorizer']
skills = model['skills']
main_model = model['resume_model']

def main():
    st.title("Resume Classification App")
    

    # Add file uploader
    uploaded_file = st.file_uploader("Upload multiple documents", type="docx", accept_multiple_files=True)
    
    results = []
    if st.button("Submit"):
        
        if uploaded_file is not None:
            for file in uploaded_file:
            
                # Read the uploaded Word document
                doc = Document(file)

                # Extract text content
                text = ' '.join([p.text for p in doc.paragraphs])

                # Extract table content
                tables = []
                for table in doc.tables:
                    for row in table.rows:
                        tables.append([cell.text for cell in row.cells])
                table_content = ' '.join([' '.join(row) for row in tables])
        
                final_text = text + " " + table_content
        
                preprocessed_text = preprocess_text(final_text)
                keyword_text = keyword_replacement(preprocessed_text)
                skill_extracted = filter_skills(keyword_text, skills)
                vectorized_text = vectorizer.transform([skill_extracted])
                experience = expDetails(text)
                if experience == None:
                    years = 'Fresher'
                else:
                    years = experience
        
        
                # Perform classification using the loaded model
                input_data = vectorized_text
                predicted_class = main_model.predict(input_data)[0]
            
                results.append({"File Name": file.name, "Job Role": predicted_class, "Experience(years)": years})

                # Display the predicted class
    #             st.write("Job Profile for: ",file.name, ":", predicted_class, " Years of Experience: ", years)
         # Display the results in a table
        if results:
            df_results = pd.DataFrame(results)
            st.table(df_results) 

if __name__ == '__main__':
    main()
